"""
5 application letter export templates — PDF and DOCX.

Each template receives:
  title          str   — document title / job title
  content        str   — the letter body (may contain markdown ## / - bullets)
  applicant_name str   — optional, from user profile
  company_name   str   — optional, extracted from application

Parsing rules (shared):
  Lines starting with "## " → sub-heading
  Lines starting with "- "  → bullet
  Blank lines               → paragraph break
  Everything else           → body paragraph
"""
from __future__ import annotations

import io
import textwrap
from datetime import datetime


def _s(text: object) -> str:
    """Latin-1 safe sanitiser. æ/ø/å/Æ/Ø/Å er i Latin-1 — konverter dem ALDRIG."""
    if not text:
        return ""
    return (
        str(text)
        .replace("—", "-").replace("–", "-").replace("•", "-")
        .replace("'", "'").replace("'", "'")
        .replace(""", '"').replace(""", '"')
        .encode("latin-1", errors="replace").decode("latin-1")
    )


def _today() -> str:
    return datetime.now().strftime("%d. %B %Y")


# ═══════════════════════════════════════════════════════════════════════════════
# PDF TEMPLATES
# ═══════════════════════════════════════════════════════════════════════════════

def _render_body_pdf(pdf, content: str, body_size: float, heading_color: tuple,
                     body_color: tuple, line_h: float = 5.5) -> None:
    """Parse markdown-ish content and render into open fpdf instance."""
    from fpdf.enums import XPos, YPos

    for raw_line in content.split("\n"):
        line = raw_line.rstrip()
        if line.startswith("## "):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", body_size + 1)
            pdf.set_text_color(*heading_color)
            pdf.cell(0, line_h + 0.5, _s(line[3:]), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_text_color(*body_color)
        elif line.startswith("# "):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", body_size + 2)
            pdf.set_text_color(*heading_color)
            pdf.cell(0, line_h + 1, _s(line[2:]), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_text_color(*body_color)
        elif line.startswith("- ") or line.startswith("* "):
            pdf.set_font("Helvetica", "", body_size)
            for wl in textwrap.wrap(line[2:], 95) or [""]:
                pdf.cell(5, line_h, "", new_x=XPos.RIGHT, new_y=YPos.TOP)
                pdf.cell(0, line_h, _s(f"- {wl}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        elif line == "" or line in ("---", "***"):
            pdf.ln(3)
        else:
            pdf.set_font("Helvetica", "", body_size)
            for wl in textwrap.wrap(line, 100) or [""]:
                pdf.cell(0, line_h, _s(wl), new_x=XPos.LMARGIN, new_y=YPos.NEXT)


# ── 1. Corporate ──────────────────────────────────────────────────────────────

def app_pdf_corporate(title: str, content: str, applicant_name: str = "",
                      company_name: str = "") -> bytes:
    """Formal professional — blue header bar, sans-serif, formal letterhead."""
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    BLUE = (30, 64, 175)
    DARK = (15, 23, 42)
    GRY = (100, 116, 139)

    pdf = FPDF()
    pdf.set_margins(22, 20, 22)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=22)

    # Header band
    pdf.set_fill_color(*BLUE)
    pdf.rect(0, 0, 210, 28, "F")
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(255, 255, 255)
    pdf.set_xy(22, 9)
    pdf.cell(166, 9, _s(title))
    if company_name:
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(180, 200, 240)
        pdf.set_xy(22, 19)
        pdf.cell(166, 5, _s(f"Til: {company_name}"))

    pdf.set_y(35)

    # Sender + date row
    if applicant_name:
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(*DARK)
        pdf.cell(0, 5, _s(applicant_name), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(0.5)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(*GRY)
    pdf.cell(0, 5, _s(_today()), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(6)

    # Body
    pdf.set_text_color(*DARK)
    _render_body_pdf(pdf, content, 10, BLUE, DARK, 5.5)

    # Footer
    pdf.set_auto_page_break(auto=False)
    pdf.set_y(-14)
    pdf.set_draw_color(*BLUE)
    pdf.set_line_width(0.4)
    pdf.line(22, pdf.get_y(), 188, pdf.get_y())
    pdf.ln(1.5)
    pdf.set_font("Helvetica", "", 7.5)
    pdf.set_text_color(*GRY)
    pdf.cell(0, 5, f"CareerOS  |  {_today()}", align="C")

    return bytes(pdf.output())


# ── 2. Executive ─────────────────────────────────────────────────────────────

def app_pdf_executive(title: str, content: str, applicant_name: str = "",
                      company_name: str = "") -> bytes:
    """Premium letterhead — Times, gold double rule, wide margins."""
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    GOLD = (140, 90, 20)
    DARK = (10, 15, 35)
    GREY = (90, 100, 120)

    pdf = FPDF()
    pdf.set_margins(30, 26, 30)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=26)

    # Letterhead
    pdf.set_font("Times", "B", 20)
    pdf.set_text_color(*DARK)
    pdf.cell(0, 10, _s(applicant_name or ""), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")

    pdf.ln(1)
    pdf.set_draw_color(*GOLD)
    pdf.set_line_width(0.8)
    pdf.line(30, pdf.get_y(), 180, pdf.get_y())
    pdf.ln(0.6)
    pdf.set_line_width(0.2)
    pdf.line(30, pdf.get_y(), 180, pdf.get_y())
    pdf.ln(5)

    # Date + addressee
    pdf.set_font("Times", "", 10)
    pdf.set_text_color(*GREY)
    pdf.cell(0, 5, _s(_today()), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="R")
    pdf.ln(3)

    if company_name:
        pdf.set_font("Times", "", 10)
        pdf.set_text_color(*GREY)
        pdf.cell(0, 5, _s(f"Att.: {company_name}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)

    # Title (RE: style)
    pdf.set_font("Times", "B", 12)
    pdf.set_text_color(*DARK)
    pdf.cell(0, 7, _s(f"Re: {title}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)

    # Body — Times, generous line height
    pdf.set_text_color(*DARK)
    for raw_line in content.split("\n"):
        line = raw_line.rstrip()
        if line.startswith("## "):
            pdf.ln(2)
            pdf.set_font("Times", "B", 11.5)
            pdf.set_text_color(*GOLD)
            pdf.cell(0, 6, _s(line[3:]), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_text_color(*DARK)
        elif line.startswith("- ") or line.startswith("* "):
            pdf.set_font("Times", "", 10.5)
            for wl in textwrap.wrap(line[2:], 88) or [""]:
                pdf.cell(0, 6, _s(f"–  {wl}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        elif line == "":
            pdf.ln(3)
        else:
            pdf.set_font("Times", "", 10.5)
            for wl in textwrap.wrap(line, 95) or [""]:
                pdf.cell(0, 6, _s(wl), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # Footer rules
    pdf.set_auto_page_break(auto=False)
    pdf.set_y(-20)
    pdf.set_draw_color(*GOLD)
    pdf.set_line_width(0.2)
    pdf.line(30, pdf.get_y(), 180, pdf.get_y())
    pdf.ln(0.6)
    pdf.set_line_width(0.8)
    pdf.line(30, pdf.get_y(), 180, pdf.get_y())
    pdf.ln(2)
    pdf.set_font("Times", "I", 8)
    pdf.set_text_color(*GREY)
    pdf.cell(0, 5, f"CareerOS  ·  {_today()}", align="C")

    return bytes(pdf.output())


# ── 3. Modern ────────────────────────────────────────────────────────────────

def app_pdf_modern(title: str, content: str, applicant_name: str = "",
                   company_name: str = "") -> bytes:
    """Clean modern — teal accent strip, ample whitespace."""
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    TEAL = (13, 148, 136)
    DARK = (15, 23, 42)
    GRY = (100, 116, 139)

    pdf = FPDF()
    pdf.set_margins(22, 22, 22)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=22)

    # Accent strip top
    pdf.set_fill_color(*TEAL)
    pdf.rect(0, 0, 210, 5, "F")
    pdf.ln(10)

    # Name + title block
    if applicant_name:
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_text_color(*DARK)
        pdf.cell(0, 7, _s(applicant_name), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(*TEAL)
    pdf.cell(0, 5, _s(title), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(*GRY)
    if company_name:
        pdf.cell(0, 5, _s(f"Til: {company_name}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.cell(0, 5, _s(_today()), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(5)

    # Teal separator
    pdf.set_draw_color(*TEAL)
    pdf.set_line_width(0.3)
    pdf.line(22, pdf.get_y(), 188, pdf.get_y())
    pdf.ln(5)

    # Body
    pdf.set_text_color(*DARK)
    _render_body_pdf(pdf, content, 10, TEAL, DARK, 5.5)

    # Bottom accent strip
    pdf.set_auto_page_break(auto=False)
    pdf.set_y(-12)
    pdf.set_fill_color(*TEAL)
    pdf.rect(0, pdf.get_y(), 210, 12, "F")
    pdf.set_font("Helvetica", "", 7.5)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 12, f"CareerOS  |  {_today()}", align="C")

    return bytes(pdf.output())


# ── 4. Technical ─────────────────────────────────────────────────────────────

def app_pdf_technical(title: str, content: str, applicant_name: str = "",
                      company_name: str = "") -> bytes:
    """Crisp minimal monochrome — wide margins, structured, no decorative colour."""
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    DARK = (10, 15, 35)
    MED = (70, 80, 100)
    LGT = (160, 170, 185)

    pdf = FPDF()
    pdf.set_margins(24, 24, 24)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=24)

    # Name line
    if applicant_name:
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(*DARK)
        pdf.cell(0, 7, _s(applicant_name), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_draw_color(*DARK)
    pdf.set_line_width(0.5)
    pdf.line(24, pdf.get_y(), 186, pdf.get_y())
    pdf.ln(3)

    # Date + RE
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(*MED)
    pdf.cell(0, 5, _s(_today()), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="R")
    pdf.ln(1)

    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(*DARK)
    re_line = f"RE: {title}"
    if company_name:
        re_line += f"  —  {company_name}"
    pdf.cell(0, 6, _s(re_line), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_draw_color(*LGT)
    pdf.set_line_width(0.2)
    pdf.line(24, pdf.get_y() + 1, 186, pdf.get_y() + 1)
    pdf.ln(6)

    # Body — strict plain, no color
    for raw_line in content.split("\n"):
        line = raw_line.rstrip()
        if line.startswith("## "):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 10.5)
            pdf.set_text_color(*DARK)
            pdf.cell(0, 6, _s(line[3:].upper()), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_draw_color(*LGT)
            pdf.set_line_width(0.2)
            pdf.line(24, pdf.get_y(), 186, pdf.get_y())
            pdf.ln(2)
        elif line.startswith("- ") or line.startswith("* "):
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(*DARK)
            for wl in textwrap.wrap(line[2:], 93) or [""]:
                pdf.cell(6, 5.5, "[+]", new_x=XPos.RIGHT, new_y=YPos.TOP)
                pdf.cell(0, 5.5, _s(wl), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        elif line == "":
            pdf.ln(3)
        else:
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(*DARK)
            for wl in textwrap.wrap(line, 100) or [""]:
                pdf.cell(0, 5.5, _s(wl), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # Footer
    pdf.set_auto_page_break(auto=False)
    pdf.set_y(-13)
    pdf.set_draw_color(*DARK)
    pdf.set_line_width(0.5)
    pdf.line(24, pdf.get_y(), 186, pdf.get_y())
    pdf.ln(1.5)
    pdf.set_font("Helvetica", "", 7.5)
    pdf.set_text_color(*MED)
    pdf.cell(0, 5, f"CareerOS  |  {_today()}", align="C")

    return bytes(pdf.output())


# ── 5. Graduate ──────────────────────────────────────────────────────────────

def app_pdf_graduate(title: str, content: str, applicant_name: str = "",
                     company_name: str = "") -> bytes:
    """Friendly professional — purple-to-blue gradient header band, readable."""
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    PUR = (79, 70, 229)
    BLUE = (59, 130, 246)
    DARK = (15, 23, 42)
    GRY = (100, 116, 139)

    pdf = FPDF()
    pdf.set_margins(20, 20, 20)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Gradient header approximation (two rects)
    pdf.set_fill_color(*PUR)
    pdf.rect(0, 0, 105, 32, "F")
    pdf.set_fill_color(*BLUE)
    pdf.rect(105, 0, 105, 32, "F")

    # Name
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(255, 255, 255)
    pdf.set_xy(20, 8)
    pdf.cell(170, 9, _s(applicant_name or title))

    # Subtitle
    pdf.set_font("Helvetica", "", 9.5)
    pdf.set_text_color(210, 215, 250)
    pdf.set_xy(20, 19)
    sub = title if applicant_name else (f"Ansøgning — {company_name}" if company_name else "")
    pdf.cell(170, 6, _s(sub))

    pdf.set_y(38)

    # Date + company
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(*GRY)
    info = _today()
    if company_name:
        info = f"{company_name}  ·  {info}"
    pdf.cell(0, 5, _s(info), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)

    # Body
    pdf.set_text_color(*DARK)
    _render_body_pdf(pdf, content, 10, PUR, DARK, 5.5)

    # Footer
    pdf.set_auto_page_break(auto=False)
    pdf.set_y(-13)
    pdf.set_font("Helvetica", "", 7.5)
    pdf.set_text_color(*GRY)
    pdf.cell(0, 5, f"CareerOS  |  {_today()}", align="C")

    return bytes(pdf.output())


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX TEMPLATES
# ═══════════════════════════════════════════════════════════════════════════════

def _rgb(tup: tuple):
    from docx.shared import RGBColor
    return RGBColor(*tup)


def _set_cell_bg(cell, rgb: tuple) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(*rgb))
    tcPr.append(shd)


def _para_space(para, before: int = 0, after: int = 0) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    pPr.append(spacing)


def _add_hr_docx(doc, color: tuple = (180, 180, 180), width: int = 6) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    para = doc.add_paragraph()
    _para_space(para, 0, 0)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(width))
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(*color))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _render_body_docx(doc, content: str, font_name: str, body_size: float,
                      heading_color: tuple, body_color: tuple) -> None:
    from docx.shared import Pt
    for raw_line in content.split("\n"):
        line = raw_line.rstrip()
        if line.startswith("## "):
            p = doc.add_paragraph()
            _para_space(p, 100, 30)
            r = p.add_run(_s(line[3:]))
            r.font.name = font_name; r.font.bold = True; r.font.size = Pt(body_size + 1)
            r.font.color.rgb = _rgb(heading_color)
        elif line.startswith("# "):
            p = doc.add_paragraph()
            _para_space(p, 100, 40)
            r = p.add_run(_s(line[2:]))
            r.font.name = font_name; r.font.bold = True; r.font.size = Pt(body_size + 2)
            r.font.color.rgb = _rgb(heading_color)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _para_space(p, 0, 30)
            r = p.add_run(_s(line[2:]))
            r.font.name = font_name; r.font.size = Pt(body_size)
            r.font.color.rgb = _rgb(body_color)
        elif line == "":
            p = doc.add_paragraph()
            _para_space(p, 0, 60)
        else:
            p = doc.add_paragraph()
            _para_space(p, 0, 40)
            r = p.add_run(_s(line))
            r.font.name = font_name; r.font.size = Pt(body_size)
            r.font.color.rgb = _rgb(body_color)


def _doc_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── 1. Corporate DOCX ────────────────────────────────────────────────────────

def app_docx_corporate(title: str, content: str, applicant_name: str = "",
                        company_name: str = "") -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt

    BLUE = (30, 64, 175)
    DARK = (15, 23, 42)
    GRY = (100, 116, 139)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0); sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)

    # Blue header table
    hdr = doc.add_table(rows=1, cols=1)
    hdr.autofit = False; hdr.columns[0].width = Inches(8.27)
    hc = hdr.rows[0].cells[0]; hc.width = Inches(8.27)
    _set_cell_bg(hc, BLUE)
    p = hc.paragraphs[0]
    p.paragraph_format.left_indent = Inches(0.35); p.paragraph_format.space_before = Pt(14)
    r = p.add_run(_s(title)); r.font.name = "Calibri"; r.font.size = Pt(16)
    r.font.bold = True; r.font.color.rgb = _rgb((255, 255, 255))
    if company_name:
        p2 = hc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.35); p2.paragraph_format.space_after = Pt(12)
        r2 = p2.add_run(_s(f"Til: {company_name}"))
        r2.font.name = "Calibri"; r2.font.size = Pt(9.5); r2.font.color.rgb = _rgb((180, 200, 240))

    doc.add_paragraph()

    if applicant_name:
        p = doc.add_paragraph(); _para_space(p, 0, 20)
        r = p.add_run(_s(applicant_name))
        r.font.name = "Calibri"; r.font.bold = True; r.font.size = Pt(11)
        r.font.color.rgb = _rgb(DARK)

    p = doc.add_paragraph(); _para_space(p, 0, 80)
    r = p.add_run(_s(_today()))
    r.font.name = "Calibri"; r.font.size = Pt(9.5); r.font.color.rgb = _rgb(GRY)

    _render_body_docx(doc, content, "Calibri", 10.5, BLUE, DARK)
    return _doc_bytes(doc)


# ── 2. Executive DOCX ────────────────────────────────────────────────────────

def app_docx_executive(title: str, content: str, applicant_name: str = "",
                        company_name: str = "") -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    GOLD = (140, 90, 20)
    DARK = (10, 15, 35)
    GREY = (90, 100, 120)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1.1); sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1.2); sec.right_margin = Inches(1.2)

    if applicant_name:
        p = doc.add_paragraph(); _para_space(p, 0, 40)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(_s(applicant_name))
        r.font.name = "Georgia"; r.font.size = Pt(20); r.font.bold = True
        r.font.color.rgb = _rgb(DARK)

    _add_hr_docx(doc, GOLD, 12)
    _add_hr_docx(doc, GOLD, 4)

    p = doc.add_paragraph(); _para_space(p, 60, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(_s(_today()))
    r.font.name = "Georgia"; r.font.size = Pt(10); r.font.italic = True
    r.font.color.rgb = _rgb(GREY)

    if company_name:
        p2 = doc.add_paragraph(); _para_space(p2, 0, 20)
        r2 = p2.add_run(_s(f"Att.: {company_name}"))
        r2.font.name = "Georgia"; r2.font.size = Pt(10); r2.font.color.rgb = _rgb(GREY)

    p3 = doc.add_paragraph(); _para_space(p3, 40, 80)
    r3 = p3.add_run(_s(f"Re: {title}"))
    r3.font.name = "Georgia"; r3.font.size = Pt(12); r3.font.bold = True
    r3.font.color.rgb = _rgb(DARK)

    _render_body_docx(doc, content, "Georgia", 10.5, GOLD, DARK)

    _add_hr_docx(doc, GOLD, 4)
    _add_hr_docx(doc, GOLD, 10)
    p = doc.add_paragraph(); _para_space(p, 40, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"CareerOS  ·  {_today()}")
    r.font.name = "Georgia"; r.font.italic = True; r.font.size = Pt(8)
    r.font.color.rgb = _rgb(GREY)
    return _doc_bytes(doc)


# ── 3. Modern DOCX ───────────────────────────────────────────────────────────

def app_docx_modern(title: str, content: str, applicant_name: str = "",
                     company_name: str = "") -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt

    TEAL = (13, 148, 136)
    TEAL_D = (10, 100, 90)
    DARK = (15, 23, 42)
    GRY = (100, 116, 139)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0); sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)

    # Teal strip header
    hdr = doc.add_table(rows=1, cols=1)
    hdr.autofit = False; hdr.columns[0].width = Inches(8.27)
    hc = hdr.rows[0].cells[0]; hc.width = Inches(8.27)
    _set_cell_bg(hc, TEAL)
    p = hc.paragraphs[0]
    p.paragraph_format.left_indent = Inches(0.35); p.paragraph_format.space_before = Pt(10)
    r = p.add_run(_s(applicant_name or title))
    r.font.name = "Calibri"; r.font.size = Pt(18); r.font.bold = True
    r.font.color.rgb = _rgb((255, 255, 255))
    if applicant_name:
        p2 = hc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.35); p2.paragraph_format.space_after = Pt(10)
        r2 = p2.add_run(_s(title))
        r2.font.name = "Calibri"; r2.font.size = Pt(10)
        r2.font.color.rgb = _rgb((200, 240, 235))

    doc.add_paragraph()

    p = doc.add_paragraph(); _para_space(p, 0, 60)
    info = _s(_today())
    if company_name: info = _s(f"{company_name}  ·  {_today()}")
    r = p.add_run(info)
    r.font.name = "Calibri"; r.font.size = Pt(9.5); r.font.color.rgb = _rgb(GRY)

    _add_hr_docx(doc, TEAL, 6)
    _render_body_docx(doc, content, "Calibri", 10.5, TEAL_D, DARK)
    return _doc_bytes(doc)


# ── 4. Technical DOCX ────────────────────────────────────────────────────────

def app_docx_technical(title: str, content: str, applicant_name: str = "",
                        company_name: str = "") -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt

    DARK = (10, 15, 35)
    MED = (70, 80, 100)
    LGT = (160, 170, 185)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1); sec.right_margin = Inches(1)

    if applicant_name:
        p = doc.add_paragraph(); _para_space(p, 0, 20)
        r = p.add_run(_s(applicant_name))
        r.font.name = "Calibri"; r.font.bold = True; r.font.size = Pt(13)
        r.font.color.rgb = _rgb(DARK)

    _add_hr_docx(doc, DARK, 8)

    p = doc.add_paragraph(); _para_space(p, 60, 0)
    re_text = f"RE: {title}"
    if company_name: re_text += f"  —  {company_name}"
    r = p.add_run(_s(re_text))
    r.font.name = "Calibri"; r.font.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = _rgb(DARK)

    p2 = doc.add_paragraph(); _para_space(p2, 0, 80)
    r2 = p2.add_run(_s(_today()))
    r2.font.name = "Calibri"; r2.font.size = Pt(9); r2.font.color.rgb = _rgb(MED)

    _add_hr_docx(doc, LGT, 3)
    _render_body_docx(doc, content, "Calibri", 10.5, DARK, DARK)
    _add_hr_docx(doc, DARK, 8)

    p = doc.add_paragraph(); _para_space(p, 30, 0)
    r = p.add_run(f"CareerOS  |  {_today()}")
    r.font.name = "Calibri"; r.font.size = Pt(8); r.font.color.rgb = _rgb(MED)
    return _doc_bytes(doc)


# ── 5. Graduate DOCX ─────────────────────────────────────────────────────────

def app_docx_graduate(title: str, content: str, applicant_name: str = "",
                       company_name: str = "") -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt

    PUR = (79, 70, 229)
    DARK = (15, 23, 42)
    GRY = (100, 116, 139)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0); sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(0.85); sec.right_margin = Inches(0.85)

    # Purple header band
    hdr = doc.add_table(rows=1, cols=1)
    hdr.autofit = False; hdr.columns[0].width = Inches(8.27)
    hc = hdr.rows[0].cells[0]; hc.width = Inches(8.27)
    _set_cell_bg(hc, PUR)
    p = hc.paragraphs[0]
    p.paragraph_format.left_indent = Inches(0.35); p.paragraph_format.space_before = Pt(14)
    r = p.add_run(_s(applicant_name or title))
    r.font.name = "Calibri"; r.font.size = Pt(18); r.font.bold = True
    r.font.color.rgb = _rgb((255, 255, 255))
    if applicant_name:
        p2 = hc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.35); p2.paragraph_format.space_after = Pt(14)
        sub = title if not company_name else f"{title}  ·  {company_name}"
        r2 = p2.add_run(_s(sub))
        r2.font.name = "Calibri"; r2.font.size = Pt(10)
        r2.font.color.rgb = _rgb((210, 215, 250))

    doc.add_paragraph()

    p = doc.add_paragraph(); _para_space(p, 0, 80)
    info = _s(_today())
    if company_name: info = _s(f"{company_name}  ·  {_today()}")
    r = p.add_run(info)
    r.font.name = "Calibri"; r.font.size = Pt(9.5); r.font.color.rgb = _rgb(GRY)

    _render_body_docx(doc, content, "Calibri", 10.5, PUR, DARK)

    p = doc.add_paragraph(); _para_space(p, 100, 0)
    r = p.add_run(f"CareerOS  |  {_today()}")
    r.font.name = "Calibri"; r.font.size = Pt(8); r.font.color.rgb = _rgb(GRY)
    return _doc_bytes(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# Routers
# ═══════════════════════════════════════════════════════════════════════════════

APP_PDF_TEMPLATES = {
    "corporate":  app_pdf_corporate,
    "executive":  app_pdf_executive,
    "modern":     app_pdf_modern,
    "technical":  app_pdf_technical,
    "graduate":   app_pdf_graduate,
}

APP_DOCX_TEMPLATES = {
    "corporate":  app_docx_corporate,
    "executive":  app_docx_executive,
    "modern":     app_docx_modern,
    "technical":  app_docx_technical,
    "graduate":   app_docx_graduate,
}


def render_app_pdf(title: str, content: str, template: str = "corporate",
                   applicant_name: str = "", company_name: str = "") -> bytes:
    fn = APP_PDF_TEMPLATES.get(template, app_pdf_corporate)
    return fn(title, content, applicant_name, company_name)


def render_app_docx(title: str, content: str, template: str = "corporate",
                    applicant_name: str = "", company_name: str = "") -> bytes:
    fn = APP_DOCX_TEMPLATES.get(template, app_docx_corporate)
    return fn(title, content, applicant_name, company_name)
