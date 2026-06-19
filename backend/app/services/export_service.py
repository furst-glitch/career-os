"""
Export Service - genererer PDF og DOCX fra CV-indhold og ansøgninger.
Bruger fpdf2 til PDF og python-docx til DOCX.
"""
from __future__ import annotations

import io
import textwrap
from datetime import datetime


def _s(text: str) -> str:
    """Sanitize text for fpdf Helvetica (Latin-1 / ISO 8859-1).
    æ/ø/å/Æ/Ø/Å ER i Latin-1 — konverter dem ALDRIG til ae/oe/aa."""
    return (
        text.replace("—", "-").replace("–", "-")
        .replace("•", "-").replace("·", "-")
        .replace("‘", "'").replace("’", "'")
        .replace("“", '"').replace("”", '"')
        .encode("latin-1", errors="replace").decode("latin-1")
    )


# ── PDF from text ─────────────────────────────────────────────────────────────

def _pdf_from_text(title: str, content: str) -> bytes:
    """Genererer velformateret PDF fra plain-text/markdown indhold."""
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    pdf = FPDF()
    pdf.set_margins(20, 20, 20)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(30, 64, 175)
    pdf.multi_cell(0, 10, _s(title), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_draw_color(148, 163, 184)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(4)

    pdf.set_text_color(30, 41, 59)
    for raw_line in content.rstrip().split("\n"):
        line = raw_line.rstrip()
        if line.startswith("## "):
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(30, 64, 175)
            pdf.multi_cell(0, 7, _s(line[3:]), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_text_color(30, 41, 59)
        elif line.startswith("# "):
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 14)
            pdf.set_text_color(30, 64, 175)
            pdf.multi_cell(0, 8, _s(line[2:]), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_text_color(30, 41, 59)
        elif line.startswith("**") and line.endswith("**"):
            pdf.set_font("Helvetica", "B", 11)
            pdf.multi_cell(0, 6, _s(line[2:-2]), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        elif line.startswith("- ") or line.startswith("* "):
            pdf.set_font("Helvetica", "", 10)
            for wl in textwrap.wrap(line[2:], width=90):
                pdf.cell(5, 5.5, "", new_x=XPos.RIGHT, new_y=YPos.TOP)
                pdf.multi_cell(0, 5.5, _s(f"- {wl}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        elif line in ("---", "***"):
            pdf.ln(2)
            pdf.set_draw_color(203, 213, 225)
            pdf.line(20, pdf.get_y(), 190, pdf.get_y())
            pdf.ln(2)
        elif line == "":
            pdf.ln(3)
        else:
            pdf.set_font("Helvetica", "", 10)
            for wl in textwrap.wrap(line, width=100) or [""]:
                pdf.multi_cell(0, 5.5, _s(wl), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_auto_page_break(auto=False)
    pdf.set_y(-15)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(148, 163, 184)
    today = datetime.now().strftime("%d/%m/%Y")
    pdf.cell(0, 10, today, align="C")

    return bytes(pdf.output())


# ── PDF from CV data ──────────────────────────────────────────────────────────

def _cv_content_to_pdf(cv_data: dict) -> bytes:
    """Struktureret PDF fra CV-data-dict."""
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    pdf = FPDF()
    pdf.set_margins(18, 18, 18)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=18)

    name = ((cv_data.get("profile") or {}).get("display_name") or "CV")
    cv_title = ((cv_data.get("master_cv") or {}).get("target_title") or "")
    summary = ((cv_data.get("master_cv") or {}).get("summary") or "")

    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(15, 23, 42)
    pdf.multi_cell(0, 10, _s(name), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    if cv_title:
        pdf.set_font("Helvetica", "", 13)
        pdf.set_text_color(37, 99, 235)
        pdf.multi_cell(0, 6, _s(cv_title), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_draw_color(37, 99, 235)
    pdf.set_line_width(0.5)
    pdf.line(18, pdf.get_y() + 2, 192, pdf.get_y() + 2)
    pdf.ln(6)

    def section_header(text: str) -> None:
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(30, 64, 175)
        pdf.multi_cell(0, 7, _s(text.upper()), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_draw_color(193, 218, 254)
        pdf.line(18, pdf.get_y(), 192, pdf.get_y())
        pdf.ln(2)
        pdf.set_text_color(15, 23, 42)

    if summary:
        section_header("Sammenfatning")
        pdf.set_font("Helvetica", "", 10)
        for wl in textwrap.wrap(summary, width=105):
            pdf.multi_cell(0, 5.5, _s(wl), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(4)

    experiences = cv_data.get("experiences") or []
    if experiences:
        section_header("Erfaring")
        for exp in experiences:
            is_cur = exp.get("is_current")
            start = (exp.get("period_start") or "")[:7]
            end = "Nu" if is_cur else (exp.get("period_end") or "")[:7]
            period = f"{start} - {end}"
            title_co = f"{exp.get('title', '')} - {exp.get('company', '')}"
            pdf.set_font("Helvetica", "B", 10)
            pdf.multi_cell(0, 5.5, _s(title_co), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("Helvetica", "I", 9)
            pdf.set_text_color(100, 116, 139)
            pdf.multi_cell(0, 4.5, _s(period), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_text_color(15, 23, 42)
            if exp.get("description"):
                pdf.set_font("Helvetica", "", 9)
                for wl in textwrap.wrap(exp["description"], width=110):
                    pdf.multi_cell(0, 4.5, _s(wl), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            for ach in (exp.get("achievements") or [])[:3]:
                pdf.set_font("Helvetica", "", 9)
                for wl in textwrap.wrap(f"- {ach}", width=108):
                    pdf.multi_cell(0, 4.5, _s(wl), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(3)

    skills = cv_data.get("skills") or []
    if skills:
        section_header("Kompetencer")
        skill_names = [s.get("name", "") for s in skills if s.get("name")]
        skill_text = "  |  ".join(skill_names)
        pdf.set_font("Helvetica", "", 10)
        for wl in textwrap.wrap(skill_text, width=100):
            pdf.multi_cell(0, 5.5, _s(wl), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(4)

    educations = cv_data.get("educations") or []
    if educations:
        section_header("Uddannelse")
        for edu in educations:
            degree_inst = f"{edu.get('degree', '')} - {edu.get('institution', '')}"
            pdf.set_font("Helvetica", "B", 10)
            pdf.multi_cell(0, 5.5, _s(degree_inst), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            ps = (edu.get("period_start") or "")[:7]
            pe = (edu.get("period_end") or "")[:7]
            pdf.set_font("Helvetica", "I", 9)
            pdf.set_text_color(100, 116, 139)
            pdf.multi_cell(0, 4.5, _s(f"{ps} - {pe}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_text_color(15, 23, 42)
            pdf.ln(3)

    certs = cv_data.get("certifications") or []
    if certs:
        section_header("Certifikater")
        for cert in certs:
            pdf.set_font("Helvetica", "", 10)
            issued = (cert.get("issued_at") or "")[:7]
            cert_line = f"- {cert.get('name', '')} - {cert.get('issuer', '')}"
            if issued:
                cert_line += f" ({issued})"
            pdf.multi_cell(0, 5.5, _s(cert_line), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(4)

    pdf.set_auto_page_break(auto=False)
    pdf.set_y(-13)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(148, 163, 184)
    today = datetime.now().strftime("%d/%m/%Y")
    pdf.cell(0, 8, today, align="C")

    return bytes(pdf.output())


# ── DOCX from text ────────────────────────────────────────────────────────────

def _docx_from_text(title: str, content: str) -> bytes:
    """Genererer DOCX fra plain-text/markdown indhold."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    h = doc.add_heading(title, 0)
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(30, 64, 175)

    for raw_line in content.split("\n"):
        line = raw_line.rstrip()
        if line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
            if p.runs:
                p.runs[0].font.color.rgb = RGBColor(30, 64, 175)
        elif line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            if p.runs:
                p.runs[0].font.color.rgb = RGBColor(30, 64, 175)
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
            run.font.size = Pt(11)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line == "---":
            doc.add_paragraph("_" * 60)
        elif line == "":
            doc.add_paragraph("")
        else:
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── DOCX from CV data ─────────────────────────────────────────────────────────

def _cv_content_to_docx(cv_data: dict) -> bytes:
    """Struktureret DOCX fra CV-data-dict."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    name = ((cv_data.get("profile") or {}).get("display_name") or "CV")
    cv_title = ((cv_data.get("master_cv") or {}).get("target_title") or "")
    summary = ((cv_data.get("master_cv") or {}).get("summary") or "")

    h = doc.add_heading(name, 0)
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(15, 23, 42)

    if cv_title:
        p = doc.add_paragraph(cv_title)
        if p.runs:
            p.runs[0].font.color.rgb = RGBColor(37, 99, 235)
            p.runs[0].font.size = Pt(13)

    if summary:
        doc.add_heading("Sammenfatning", level=2)
        doc.add_paragraph(summary)

    experiences = cv_data.get("experiences") or []
    if experiences:
        doc.add_heading("Erfaring", level=2)
        for exp in experiences:
            is_cur = exp.get("is_current")
            start = (exp.get("period_start") or "")[:7]
            end = "Nu" if is_cur else (exp.get("period_end") or "")[:7]
            p = doc.add_paragraph()
            run = p.add_run(f"{exp.get('title', '')} - {exp.get('company', '')}")
            run.bold = True
            doc.add_paragraph(f"{start} - {end}")
            if exp.get("description"):
                doc.add_paragraph(exp["description"])
            for ach in (exp.get("achievements") or [])[:3]:
                doc.add_paragraph(ach, style="List Bullet")

    skills = cv_data.get("skills") or []
    if skills:
        doc.add_heading("Kompetencer", level=2)
        skill_names = [s.get("name", "") for s in skills if s.get("name")]
        doc.add_paragraph("  |  ".join(skill_names))

    educations = cv_data.get("educations") or []
    if educations:
        doc.add_heading("Uddannelse", level=2)
        for edu in educations:
            p = doc.add_paragraph()
            run = p.add_run(f"{edu.get('degree', '')} - {edu.get('institution', '')}")
            run.bold = True

    certs = cv_data.get("certifications") or []
    if certs:
        doc.add_heading("Certifikater", level=2)
        for cert in certs:
            issued = (cert.get("issued_at") or "")[:7]
            cert_line = f"{cert.get('name', '')} - {cert.get('issuer', '')}"
            if issued:
                cert_line += f" ({issued})"
            doc.add_paragraph(cert_line, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Generated CV renderer (tekst + kontakthoved) ─────────────────────────────

def _generated_cv_to_pdf(title: str, content: str, profile: dict) -> bytes:
    """Renderer et AI-genereret CV (plain text/markdown) med kontakthoved."""
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    BLUE = (30, 64, 175)
    DARK = (15, 23, 42)
    GRY = (100, 116, 139)
    LGT = (148, 163, 184)

    pdf = FPDF()
    pdf.set_margins(20, 20, 20)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Navn
    name = _s(
        profile.get("full_name") or profile.get("display_name") or ""
    )
    if name:
        pdf.set_font("Helvetica", "B", 20)
        pdf.set_text_color(*DARK)
        pdf.multi_cell(0, 10, name, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # Kontaktlinje: email · tlf · lokation
    contact_parts = []
    if profile.get("email"):
        contact_parts.append(_s(profile["email"]))
    if profile.get("phone"):
        contact_parts.append(_s(profile["phone"]))
    if profile.get("location"):
        contact_parts.append(_s(profile["location"]))
    if contact_parts:
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*GRY)
        pdf.multi_cell(0, 5, "  ·  ".join(contact_parts), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    if profile.get("linkedin_url"):
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*GRY)
        pdf.multi_cell(0, 5, _s(profile["linkedin_url"]), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # Skillelinje
    pdf.ln(2)
    pdf.set_draw_color(*BLUE)
    pdf.set_line_width(0.5)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(5)

    # Indhold — renderer markdown
    pdf.set_text_color(*DARK)
    for raw_line in content.rstrip().split("\n"):
        line = raw_line.rstrip()
        if line.startswith("## "):
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(*BLUE)
            pdf.multi_cell(0, 6, _s(line[3:]), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_draw_color(*LGT)
            pdf.set_line_width(0.2)
            pdf.line(20, pdf.get_y(), 190, pdf.get_y())
            pdf.ln(2)
            pdf.set_text_color(*DARK)
        elif line.startswith("# "):
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_text_color(*DARK)
            pdf.multi_cell(0, 7, _s(line[2:]), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_text_color(*DARK)
        elif line.startswith("**") and line.endswith("**"):
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(*DARK)
            pdf.multi_cell(0, 5.5, _s(line[2:-2]), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        elif line.startswith("- ") or line.startswith("* "):
            pdf.set_font("Helvetica", "", 9.5)
            pdf.set_text_color(*DARK)
            for wl in textwrap.wrap(line[2:], width=100) or [""]:
                pdf.cell(5, 5, "", new_x=XPos.RIGHT, new_y=YPos.TOP)
                pdf.multi_cell(0, 5, _s(f"- {wl}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        elif line in ("---", "***"):
            pdf.ln(1)
            pdf.set_draw_color(*LGT)
            pdf.set_line_width(0.2)
            pdf.line(20, pdf.get_y(), 190, pdf.get_y())
            pdf.ln(2)
        elif line == "":
            pdf.ln(3)
        else:
            pdf.set_font("Helvetica", "", 9.5)
            pdf.set_text_color(*DARK)
            for wl in textwrap.wrap(line, width=105) or [""]:
                pdf.multi_cell(0, 5, _s(wl), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_auto_page_break(auto=False)
    pdf.set_y(-13)
    pdf.set_font("Helvetica", "I", 7.5)
    pdf.set_text_color(*LGT)
    today = datetime.now().strftime("%d/%m/%Y")
    pdf.cell(0, 5, today, align="C")

    return bytes(pdf.output())


def export_generated_cv_as_pdf(title: str, content: str, profile: dict) -> bytes:
    return _generated_cv_to_pdf(title, content, profile)


def export_generated_cv_as_docx(title: str, content: str, profile: dict) -> bytes:
    """DOCX-version af genereret CV med kontakthoved."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    name = profile.get("full_name") or profile.get("display_name") or ""
    if name:
        h = doc.add_heading(name, 0)
        if h.runs:
            h.runs[0].font.color.rgb = RGBColor(15, 23, 42)

    contact_parts = [
        v for v in [
            profile.get("email"), profile.get("phone"),
            profile.get("location"), profile.get("linkedin_url"),
        ] if v
    ]
    if contact_parts:
        p = doc.add_paragraph("  ·  ".join(contact_parts))
        if p.runs:
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph("")

    for raw_line in content.split("\n"):
        line = raw_line.rstrip()
        if line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
            if p.runs:
                p.runs[0].font.color.rgb = RGBColor(30, 64, 175)
        elif line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            if p.runs:
                p.runs[0].font.color.rgb = RGBColor(30, 64, 175)
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
            run.font.size = Pt(10)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line == "":
            doc.add_paragraph("")
        else:
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Public interface ──────────────────────────────────────────────────────────

def export_text_as_pdf(
    title: str,
    content: str,
    template: str = "corporate",
    applicant_name: str = "",
    company_name: str = "",
    profile: dict | None = None,
) -> bytes:
    from app.services.app_export_templates import render_app_pdf
    name = (profile or {}).get("full_name") or (profile or {}).get("display_name") or applicant_name
    return render_app_pdf(title, content, template, name, company_name, profile or {})


def export_text_as_docx(
    title: str,
    content: str,
    template: str = "corporate",
    applicant_name: str = "",
    company_name: str = "",
    profile: dict | None = None,
) -> bytes:
    from app.services.app_export_templates import render_app_docx
    name = (profile or {}).get("full_name") or (profile or {}).get("display_name") or applicant_name
    return render_app_docx(title, content, template, name, company_name, profile or {})


def export_cv_as_pdf(cv_data: dict, template: str = "ats_professional") -> bytes:
    from app.services.cv_pdf_templates import render_cv_pdf
    return render_cv_pdf(cv_data, template)


def export_cv_as_docx(cv_data: dict, template: str = "ats_professional") -> bytes:
    from app.services.cv_docx_templates import render_cv_docx
    return render_cv_docx(cv_data, template)
