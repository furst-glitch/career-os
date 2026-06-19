"""
5 CV DOCX templates using python-docx.
Each function accepts a cv_data dict and returns DOCX bytes.
"""
from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement


# ── Helpers ───────────────────────────────────────────────────────────────────

def _s(text: object) -> str:
    return str(text) if text else ""


def _period(exp: dict) -> str:
    start = (_s(exp.get("period_start")))[:7]
    end = "Nu" if exp.get("is_current") else (_s(exp.get("period_end")))[:7]
    return f"{start} – {end}" if start else ""


def _today() -> str:
    return datetime.now().strftime("%d/%m/%Y")


def _rgb(hex_or_tuple) -> RGBColor:
    if isinstance(hex_or_tuple, tuple):
        return RGBColor(*hex_or_tuple)
    h = hex_or_tuple.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_bg(cell, rgb: tuple) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    hex_color = "%02X%02X%02X" % rgb
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _para_space(para, before: int = 0, after: int = 0) -> None:
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    pPr.append(spacing)


def _add_hr(doc: Document, color: tuple = (180, 180, 180), width: int = 6) -> None:
    para = doc.add_paragraph()
    _para_space(para, 0, 0)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(width))
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "%02X%02X%02X" % color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_margins(doc: Document, top: float, bottom: float, left: float, right: float) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def _doc_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── 1. ATS Professional ───────────────────────────────────────────────────────

def cv_docx_ats(cv_data: dict) -> bytes:
    doc = Document()
    _set_margins(doc, 0.9, 0.9, 0.85, 0.85)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    profile = cv_data.get("profile") or {}
    master = cv_data.get("master_cv") or {}
    name = _s(profile.get("display_name"))
    title_str = _s(master.get("target_title"))
    summary = _s(master.get("summary"))

    # Name
    p = doc.add_paragraph()
    _para_space(p, 0, 40)
    r = p.add_run(name)
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = _rgb((10, 15, 35))

    if title_str:
        p2 = doc.add_paragraph()
        _para_space(p2, 0, 60)
        r2 = p2.add_run(title_str)
        r2.font.size = Pt(11)
        r2.font.color.rgb = _rgb((70, 85, 105))

    def section_head(label: str) -> None:
        _add_hr(doc, (10, 15, 35), 8)
        p = doc.add_paragraph()
        _para_space(p, 60, 40)
        r = p.add_run(label.upper())
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = _rgb((10, 15, 35))

    def body(text: str, size: float = 10) -> None:
        p = doc.add_paragraph()
        _para_space(p, 0, 30)
        r = p.add_run(_s(text))
        r.font.size = Pt(size)

    if summary:
        section_head("Professional Summary")
        body(summary)

    exps = cv_data.get("experiences") or []
    if exps:
        section_head("Experience")
        for exp in exps:
            p = doc.add_paragraph()
            _para_space(p, 40, 0)
            r = p.add_run(f"{exp.get('title','')} | {exp.get('company','')}")
            r.font.bold = True
            r.font.size = Pt(10)
            period = _period(exp)
            if period:
                p2 = doc.add_paragraph()
                _para_space(p2, 0, 20)
                r2 = p2.add_run(period)
                r2.font.size = Pt(9)
                r2.font.color.rgb = _rgb((90, 100, 120))
            if exp.get("description"):
                body(exp["description"])
            for ach in (exp.get("achievements") or [])[:4]:
                p3 = doc.add_paragraph(style="List Bullet")
                _para_space(p3, 0, 20)
                r3 = p3.add_run(_s(ach))
                r3.font.size = Pt(9.5)

    skills = cv_data.get("skills") or []
    if skills:
        section_head("Skills")
        names = " | ".join(_s(s.get("name")) for s in skills if s.get("name"))
        body(names)

    edus = cv_data.get("educations") or []
    if edus:
        section_head("Education")
        for edu in edus:
            p = doc.add_paragraph()
            _para_space(p, 40, 0)
            r = p.add_run(f"{edu.get('degree','')} – {edu.get('institution','')}")
            r.font.bold = True
            r.font.size = Pt(10)
            ps = (_s(edu.get("period_start")))[:7]
            pe = (_s(edu.get("period_end")))[:7]
            if ps:
                p2 = doc.add_paragraph()
                _para_space(p2, 0, 30)
                r2 = p2.add_run(f"{ps} – {pe}")
                r2.font.size = Pt(9)
                r2.font.color.rgb = _rgb((90, 100, 120))

    certs = cv_data.get("certifications") or []
    if certs:
        section_head("Certifications")
        for cert in certs:
            issued = (_s(cert.get("issued_at")))[:7]
            line = f"{cert.get('name','')} – {cert.get('issuer','')}"
            if issued:
                line += f" ({issued})"
            body(f"• {line}")

    return _doc_bytes(doc)


# ── 2. Modern Professional ────────────────────────────────────────────────────

def cv_docx_modern(cv_data: dict) -> bytes:
    doc = Document()
    _set_margins(doc, 0, 0, 0, 0)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    NAV = (15, 40, 80)
    ACC = (59, 130, 246)
    WHT = (255, 255, 255)
    DRK = (15, 23, 42)
    GRY = (100, 116, 139)

    profile = cv_data.get("profile") or {}
    master = cv_data.get("master_cv") or {}
    name = _s(profile.get("display_name"))
    title_str = _s(master.get("target_title"))
    summary = _s(master.get("summary"))

    # Full-width navy header row via 1-col table
    hdr_tbl = doc.add_table(rows=1, cols=1)
    hdr_tbl.autofit = False
    hdr_tbl.columns[0].width = Inches(8.27)
    hdr_cell = hdr_tbl.rows[0].cells[0]
    _set_cell_bg(hdr_cell, NAV)
    hdr_cell.width = Inches(8.27)

    p = hdr_cell.paragraphs[0]
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run(name)
    r.font.name = "Calibri"
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = _rgb(WHT)

    if title_str:
        p2 = hdr_cell.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.35)
        p2.paragraph_format.space_after = Pt(12)
        r2 = p2.add_run(title_str)
        r2.font.name = "Calibri"
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = _rgb(ACC)

    # Two-column body: sidebar | main
    SB_W = Inches(2.4)
    MAIN_W = Inches(5.0)

    doc.add_paragraph()  # spacer

    def section_head_main(label: str, para_list: list, color: tuple = ACC) -> None:
        p = doc.add_paragraph()
        _para_space(p, 100, 40)
        p.paragraph_format.left_indent = Inches(2.7)
        r = p.add_run(label.upper())
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = _rgb(color)
        _add_hr(doc, color, 4)

    def main_body(text: str, indent: float = 2.7) -> None:
        p = doc.add_paragraph()
        _para_space(p, 0, 30)
        p.paragraph_format.left_indent = Inches(indent)
        r = p.add_run(_s(text))
        r.font.size = Pt(9.5)
        r.font.color.rgb = _rgb(DRK)

    if summary:
        section_head_main("Profile", [])
        main_body(summary)

    exps = cv_data.get("experiences") or []
    if exps:
        section_head_main("Experience", [])
        for exp in exps:
            p = doc.add_paragraph()
            _para_space(p, 50, 0)
            p.paragraph_format.left_indent = Inches(2.7)
            r = p.add_run(_s(exp.get("title", "")))
            r.font.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = _rgb(DRK)

            p2 = doc.add_paragraph()
            _para_space(p2, 0, 30)
            p2.paragraph_format.left_indent = Inches(2.7)
            r2 = p2.add_run(f"{_s(exp.get('company',''))}   {_period(exp)}")
            r2.font.size = Pt(9); r2.font.color.rgb = _rgb(GRY)

            if exp.get("description"):
                main_body(exp["description"])
            for ach in (exp.get("achievements") or [])[:3]:
                main_body(f"• {ach}")

    skills = cv_data.get("skills") or []
    if skills:
        p = doc.add_paragraph()
        _para_space(p, 80, 40)
        p.paragraph_format.left_indent = Inches(0.3)
        r = p.add_run("SKILLS")
        r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = _rgb(ACC)
        for sk in skills[:14]:
            p2 = doc.add_paragraph()
            _para_space(p2, 0, 20)
            p2.paragraph_format.left_indent = Inches(0.3)
            r2 = p2.add_run(f"• {_s(sk.get('name',''))}")
            r2.font.size = Pt(9); r2.font.color.rgb = _rgb(WHT)

    edus = cv_data.get("educations") or []
    if edus:
        section_head_main("Education", [])
        for edu in edus:
            main_body(f"{edu.get('degree','')} – {edu.get('institution','')}")
            ps = (_s(edu.get("period_start")))[:7]; pe = (_s(edu.get("period_end")))[:7]
            if ps:
                main_body(f"{ps} – {pe}")

    return _doc_bytes(doc)


# ── 3. Executive ─────────────────────────────────────────────────────────────

def cv_docx_executive(cv_data: dict) -> bytes:
    doc = Document()
    _set_margins(doc, 1.1, 1.0, 1.1, 1.1)
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(10.5)

    GOLD = (160, 100, 30)
    DARK = (10, 15, 35)
    GREY = (90, 100, 120)

    profile = cv_data.get("profile") or {}
    master = cv_data.get("master_cv") or {}
    name = _s(profile.get("display_name"))
    title_str = _s(master.get("target_title"))
    summary = _s(master.get("summary"))

    # Centered name
    p = doc.add_paragraph()
    _para_space(p, 0, 40)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(name)
    r.font.name = "Georgia"; r.font.size = Pt(24); r.font.bold = True
    r.font.color.rgb = _rgb(DARK)

    if title_str:
        p2 = doc.add_paragraph()
        _para_space(p2, 0, 80)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(title_str)
        r2.font.name = "Georgia"; r2.font.size = Pt(12); r2.font.italic = True
        r2.font.color.rgb = _rgb(GREY)

    _add_hr(doc, GOLD, 12)
    _add_hr(doc, GOLD, 4)

    def section(label: str) -> None:
        _add_hr(doc, GOLD, 6)
        p = doc.add_paragraph()
        _para_space(p, 80, 40)
        r = p.add_run(label.upper())
        r.font.name = "Georgia"; r.font.bold = True; r.font.size = Pt(11)
        r.font.color.rgb = _rgb(GOLD)

    def body(text: str, size: float = 10.5) -> None:
        p = doc.add_paragraph()
        _para_space(p, 0, 40)
        r = p.add_run(_s(text))
        r.font.name = "Georgia"; r.font.size = Pt(size)
        r.font.color.rgb = _rgb(DARK)

    if summary:
        section("Executive Summary")
        body(summary)

    exps = cv_data.get("experiences") or []
    if exps:
        section("Professional Experience")
        for exp in exps:
            p = doc.add_paragraph()
            _para_space(p, 60, 0)
            r = p.add_run(f"{exp.get('title','')}  ·  {exp.get('company','')}")
            r.font.name = "Georgia"; r.font.bold = True; r.font.size = Pt(11)
            r.font.color.rgb = _rgb(DARK)
            period = _period(exp)
            if period:
                p2 = doc.add_paragraph()
                _para_space(p2, 0, 30)
                r2 = p2.add_run(period)
                r2.font.name = "Georgia"; r2.font.italic = True; r2.font.size = Pt(9.5)
                r2.font.color.rgb = _rgb(GREY)
            if exp.get("description"):
                body(exp["description"])
            for ach in (exp.get("achievements") or [])[:4]:
                body(f"–  {ach}", 10)

    skills = cv_data.get("skills") or []
    if skills:
        section("Core Competencies")
        names = "  ·  ".join(_s(s.get("name")) for s in skills if s.get("name"))
        body(names)

    edus = cv_data.get("educations") or []
    if edus:
        section("Education")
        for edu in edus:
            body(f"{edu.get('degree','')}  ·  {edu.get('institution','')}", 10.5)
            ps = (_s(edu.get("period_start")))[:7]; pe = (_s(edu.get("period_end")))[:7]
            if ps:
                p = doc.add_paragraph()
                _para_space(p, 0, 40)
                r = p.add_run(f"{ps} – {pe}")
                r.font.name = "Georgia"; r.font.italic = True; r.font.size = Pt(9.5)
                r.font.color.rgb = _rgb(GREY)

    certs = cv_data.get("certifications") or []
    if certs:
        section("Certifications")
        for cert in certs:
            issued = (_s(cert.get("issued_at")))[:7]
            line = f"–  {cert.get('name','')}  |  {cert.get('issuer','')}"
            if issued: line += f"  ({issued})"
            body(line, 10)

    _add_hr(doc, GOLD, 4)
    _add_hr(doc, GOLD, 10)
    p = doc.add_paragraph()
    _para_space(p, 40, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"CareerOS  ·  {_today()}")
    r.font.name = "Georgia"; r.font.italic = True; r.font.size = Pt(8)
    r.font.color.rgb = _rgb(GREY)

    return _doc_bytes(doc)


# ── 4. Minimal Nordic ─────────────────────────────────────────────────────────

def cv_docx_nordic(cv_data: dict) -> bytes:
    doc = Document()
    _set_margins(doc, 1.1, 1.1, 1.0, 1.0)
    style = doc.styles["Normal"]
    style.font.name = "Calibri Light"
    style.font.size = Pt(10.5)

    DARK = (25, 35, 50)
    MED = (80, 95, 115)
    LGT = (180, 195, 210)

    profile = cv_data.get("profile") or {}
    master = cv_data.get("master_cv") or {}
    name = _s(profile.get("display_name"))
    title_str = _s(master.get("target_title"))
    summary = _s(master.get("summary"))

    # Name — generous size, light weight
    p = doc.add_paragraph()
    _para_space(p, 0, 40)
    r = p.add_run(name)
    r.font.name = "Calibri Light"; r.font.size = Pt(28)
    r.font.color.rgb = _rgb(DARK)

    if title_str:
        p2 = doc.add_paragraph()
        _para_space(p2, 0, 100)
        r2 = p2.add_run(title_str)
        r2.font.name = "Calibri Light"; r2.font.size = Pt(12)
        r2.font.color.rgb = _rgb(MED)

    _add_hr(doc, LGT, 4)

    def section(label: str) -> None:
        p = doc.add_paragraph()
        _para_space(p, 140, 40)
        r = p.add_run(label.upper())
        r.font.name = "Calibri Light"; r.font.size = Pt(8.5)
        r.font.color.rgb = _rgb(MED)
        _add_hr(doc, LGT, 3)

    def body(text: str, size: float = 10.5) -> None:
        p = doc.add_paragraph()
        _para_space(p, 0, 50)
        r = p.add_run(_s(text))
        r.font.name = "Calibri Light"; r.font.size = Pt(size)
        r.font.color.rgb = _rgb(DARK)

    if summary:
        section("Summary")
        body(summary)

    exps = cv_data.get("experiences") or []
    if exps:
        section("Experience")
        for exp in exps:
            p = doc.add_paragraph()
            _para_space(p, 80, 0)
            r = p.add_run(_s(exp.get("title", "")))
            r.font.name = "Calibri Light"; r.font.bold = True; r.font.size = Pt(11)
            r.font.color.rgb = _rgb(DARK)

            p2 = doc.add_paragraph()
            _para_space(p2, 0, 20)
            r_co = p2.add_run(_s(exp.get("company", "")))
            r_co.font.name = "Calibri Light"; r_co.font.size = Pt(10)
            r_co.font.color.rgb = _rgb(MED)
            if _period(exp):
                r_sep = p2.add_run(f"   {_period(exp)}")
                r_sep.font.name = "Calibri Light"; r_sep.font.size = Pt(9)
                r_sep.font.color.rgb = _rgb(LGT)

            if exp.get("description"):
                body(exp["description"], 10.5)
            for ach in (exp.get("achievements") or [])[:3]:
                p3 = doc.add_paragraph()
                _para_space(p3, 0, 30)
                r3 = p3.add_run(f"— {_s(ach)}")
                r3.font.name = "Calibri Light"; r3.font.size = Pt(10)
                r3.font.color.rgb = _rgb(MED)

    skills = cv_data.get("skills") or []
    if skills:
        section("Skills")
        names = "   ·   ".join(_s(s.get("name")) for s in skills if s.get("name"))
        body(names, 10)

    edus = cv_data.get("educations") or []
    if edus:
        section("Education")
        for edu in edus:
            p = doc.add_paragraph()
            _para_space(p, 80, 0)
            r = p.add_run(_s(edu.get("degree", "")))
            r.font.name = "Calibri Light"; r.font.bold = True; r.font.size = Pt(11)
            r.font.color.rgb = _rgb(DARK)
            p2 = doc.add_paragraph()
            _para_space(p2, 0, 30)
            r2 = p2.add_run(_s(edu.get("institution", "")))
            r2.font.name = "Calibri Light"; r2.font.size = Pt(10)
            r2.font.color.rgb = _rgb(MED)
            ps = (_s(edu.get("period_start")))[:7]; pe = (_s(edu.get("period_end")))[:7]
            if ps:
                p3 = doc.add_paragraph()
                _para_space(p3, 0, 50)
                r3 = p3.add_run(f"{ps} – {pe}")
                r3.font.name = "Calibri Light"; r3.font.size = Pt(9)
                r3.font.color.rgb = _rgb(LGT)

    certs = cv_data.get("certifications") or []
    if certs:
        section("Certifications")
        for cert in certs:
            issued = (_s(cert.get("issued_at")))[:7]
            line = f"{cert.get('name','')}  ·  {cert.get('issuer','')}"
            if issued: line += f"  ·  {issued}"
            body(line, 10)

    _add_hr(doc, LGT, 3)
    p = doc.add_paragraph()
    _para_space(p, 30, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"CareerOS  ·  {_today()}")
    r.font.name = "Calibri Light"; r.font.size = Pt(7.5)
    r.font.color.rgb = _rgb(LGT)

    return _doc_bytes(doc)


# ── 5. Creative Professional ──────────────────────────────────────────────────

def cv_docx_creative(cv_data: dict) -> bytes:
    doc = Document()
    _set_margins(doc, 0, 0.9, 0.85, 0.85)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    TEAL = (13, 148, 136)
    TEAL_D = (10, 100, 90)
    DARK = (15, 23, 42)
    GRY = (100, 116, 139)
    WHT = (255, 255, 255)
    TEAL_L = (204, 240, 237)

    profile = cv_data.get("profile") or {}
    master = cv_data.get("master_cv") or {}
    name = _s(profile.get("display_name"))
    title_str = _s(master.get("target_title"))
    summary = _s(master.get("summary"))

    # Teal header band via table
    hdr = doc.add_table(rows=1, cols=1)
    hdr.autofit = False
    hdr.columns[0].width = Inches(8.27)
    hc = hdr.rows[0].cells[0]
    hc.width = Inches(8.27)
    _set_cell_bg(hc, TEAL)
    p = hc.paragraphs[0]
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(name)
    r.font.name = "Calibri"; r.font.size = Pt(22); r.font.bold = True
    r.font.color.rgb = _rgb(WHT)
    if title_str:
        p2 = hc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.35)
        p2.paragraph_format.space_after = Pt(14)
        r2 = p2.add_run(title_str)
        r2.font.name = "Calibri"; r2.font.size = Pt(11)
        r2.font.color.rgb = _rgb(TEAL_L)

    # Thin accent strip table
    strip = doc.add_table(rows=1, cols=1)
    strip.autofit = False
    strip.columns[0].width = Inches(8.27)
    sc = strip.rows[0].cells[0]
    sc.width = Inches(8.27)
    _set_cell_bg(sc, TEAL_D)
    sc.paragraphs[0].paragraph_format.space_before = Pt(2)
    sc.paragraphs[0].paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    def section(label: str) -> None:
        p = doc.add_paragraph()
        _para_space(p, 80, 40)
        r = p.add_run(label.upper())
        r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = _rgb(TEAL)
        _add_hr(doc, TEAL, 6)

    def body(text: str, size: float = 9.5, color: tuple = None) -> None:
        p = doc.add_paragraph()
        _para_space(p, 0, 40)
        r = p.add_run(_s(text))
        r.font.size = Pt(size)
        r.font.color.rgb = _rgb(color or DARK)

    if summary:
        section("Profile")
        body(summary)

    exps = cv_data.get("experiences") or []
    if exps:
        section("Experience")
        for exp in exps:
            p = doc.add_paragraph()
            _para_space(p, 60, 0)
            r = p.add_run(_s(exp.get("title", "")))
            r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = _rgb(DARK)

            p2 = doc.add_paragraph()
            _para_space(p2, 0, 30)
            r_co = p2.add_run(_s(exp.get("company", "")))
            r_co.font.size = Pt(9.5); r_co.font.color.rgb = _rgb(GRY)
            if _period(exp):
                r_p = p2.add_run(f"   {_period(exp)}")
                r_p.font.size = Pt(9); r_p.font.color.rgb = _rgb(GRY)

            if exp.get("description"):
                body(exp["description"])
            for ach in (exp.get("achievements") or [])[:3]:
                p3 = doc.add_paragraph()
                _para_space(p3, 0, 25)
                r_bullet = p3.add_run("› ")
                r_bullet.font.size = Pt(9.5); r_bullet.font.color.rgb = _rgb(TEAL)
                r3 = p3.add_run(_s(ach))
                r3.font.size = Pt(9.5); r3.font.color.rgb = _rgb(DARK)

    skills = cv_data.get("skills") or []
    if skills:
        section("Skills")
        names = _s(  "  ·  ".join(_s(s.get("name")) for s in skills if s.get("name")))
        body(names, 9.5, TEAL_D)

    edus = cv_data.get("educations") or []
    if edus:
        section("Education")
        for edu in edus:
            body(f"{edu.get('degree','')} — {edu.get('institution','')}", 10)
            ps = (_s(edu.get("period_start")))[:7]; pe = (_s(edu.get("period_end")))[:7]
            if ps:
                body(f"{ps} – {pe}", 9, GRY)

    certs = cv_data.get("certifications") or []
    if certs:
        section("Certifications")
        for cert in certs:
            issued = (_s(cert.get("issued_at")))[:7]
            line = f"{cert.get('name','')} – {cert.get('issuer','')}"
            if issued: line += f" ({issued})"
            body(line)

    p = doc.add_paragraph()
    _para_space(p, 100, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"CareerOS  |  {_today()}")
    r.font.size = Pt(7.5); r.font.color.rgb = _rgb(GRY)

    return _doc_bytes(doc)


# ── Router ────────────────────────────────────────────────────────────────────

CV_DOCX_TEMPLATES: dict[str, object] = {
    "ats_professional":    cv_docx_ats,
    "modern_professional": cv_docx_modern,
    "executive":           cv_docx_executive,
    "minimal_nordic":      cv_docx_nordic,
    "creative_professional": cv_docx_creative,
}


def render_cv_docx(cv_data: dict, template: str = "ats_professional") -> bytes:
    fn = CV_DOCX_TEMPLATES.get(template, cv_docx_ats)
    return fn(cv_data)  # type: ignore[operator]
